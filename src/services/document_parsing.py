import re
import shutil
import subprocess
import tempfile
import time
from dataclasses import dataclass
from pathlib import Path
from typing import Protocol

import anyio

from src.core.settings import Settings, get_settings

PARSER_VERSION = "lightweight-v2"
PROCESSABLE_EXTENSIONS = {"pdf", "docx", "txt", "md", "rtf"}


class UnsupportedDocumentTypeError(Exception):
    pass


class DocumentNeedsOcrError(Exception):
    pass


class DocumentOcrUnavailableError(DocumentNeedsOcrError):
    pass


class DocumentOcrTimeoutError(Exception):
    pass


class DocumentOcrFailedError(Exception):
    pass


@dataclass(frozen=True)
class ParsedPage:
    page_number: int
    text: str
    char_start: int
    char_end: int


@dataclass(frozen=True)
class ParsedSection:
    title: str
    page_start: int | None
    char_start: int
    char_end: int


@dataclass(frozen=True)
class ParsedDocumentContent:
    filename: str
    file_extension: str
    text: str
    pages: list[ParsedPage]
    sections: list[ParsedSection]
    parser_version: str
    metadata: dict


@dataclass(frozen=True)
class OcrTextResult:
    text: str
    metadata: dict


class DocumentOcrService(Protocol):
    def extract_text(self, *, content: bytes, page_count: int) -> OcrTextResult: ...


class OcrmypdfTesseractService:
    def __init__(self, settings: Settings) -> None:
        self.settings = settings

    def extract_text(self, *, content: bytes, page_count: int) -> OcrTextResult:
        if page_count > self.settings.ocr_max_pages:
            raise DocumentOcrFailedError(
                f"PDF has {page_count} pages, which exceeds OCR_MAX_PAGES={self.settings.ocr_max_pages}."
            )
        if shutil.which("ocrmypdf") is None:
            raise DocumentOcrUnavailableError("OCRmyPDF is not installed in this runtime.")
        started = time.perf_counter()
        with tempfile.TemporaryDirectory(prefix="buro-ocr-") as temp_dir:
            temp_path = Path(temp_dir)
            input_path = temp_path / "input.pdf"
            output_path = temp_path / "output.pdf"
            sidecar_path = temp_path / "sidecar.txt"
            input_path.write_bytes(content)
            command = [
                "ocrmypdf",
                "--force-ocr",
                "--sidecar",
                str(sidecar_path),
                "--language",
                self.settings.ocr_languages,
                "--jobs",
                "1",
                str(input_path),
                str(output_path),
            ]
            try:
                completed = subprocess.run(
                    command,
                    check=False,
                    capture_output=True,
                    text=True,
                    timeout=self.settings.ocr_timeout_seconds,
                )
            except subprocess.TimeoutExpired as exc:
                raise DocumentOcrTimeoutError("OCR timed out.") from exc
            if completed.returncode != 0:
                stderr = completed.stderr.strip()[:1000]
                raise DocumentOcrFailedError(stderr or "OCR failed.")
            text = sidecar_path.read_text(encoding="utf-8", errors="ignore")
        return OcrTextResult(
            text=text,
            metadata={
                "ocr_engine": "ocrmypdf+tesseract",
                "ocr_languages": self.settings.ocr_languages,
                "ocr_page_count": page_count,
                "ocr_duration_ms": round((time.perf_counter() - started) * 1000),
                "ocr_timeout": False,
            },
        )


class DocumentParser:
    def __init__(
        self,
        *,
        settings: Settings | None = None,
        ocr_service: DocumentOcrService | None = None,
    ) -> None:
        self.settings = settings or get_settings()
        self.ocr_service = ocr_service or OcrmypdfTesseractService(self.settings)

    async def parse(
        self,
        *,
        filename: str,
        file_extension: str,
        content: bytes,
    ) -> ParsedDocumentContent:
        normalized_extension = file_extension.lower().removeprefix(".")
        if normalized_extension not in PROCESSABLE_EXTENSIONS:
            raise UnsupportedDocumentTypeError(
                f"File extension '{normalized_extension}' is not supported for processing."
            )

        if normalized_extension == "pdf":
            return await anyio.to_thread.run_sync(
                self._parse_pdf,
                filename,
                normalized_extension,
                content,
            )
        if normalized_extension == "docx":
            return await anyio.to_thread.run_sync(
                self._parse_docx,
                filename,
                normalized_extension,
                content,
            )
        if normalized_extension == "rtf":
            return await anyio.to_thread.run_sync(
                self._parse_rtf,
                filename,
                normalized_extension,
                content,
            )

        text = self._decode_text(content)
        return self.parse_plain_text(
            filename=filename,
            file_extension=normalized_extension,
            text=text,
        )

    @staticmethod
    def parse_plain_text(
        *,
        filename: str,
        file_extension: str,
        text: str,
    ) -> ParsedDocumentContent:
        normalized_text = normalize_extracted_text(text)
        if not normalized_text:
            raise DocumentNeedsOcrError("Document does not contain extractable text.")

        page = ParsedPage(
            page_number=1,
            text=normalized_text,
            char_start=0,
            char_end=len(normalized_text),
        )
        return ParsedDocumentContent(
            filename=filename,
            file_extension=file_extension,
            text=normalized_text,
            pages=[page],
            sections=detect_sections(normalized_text, pages=[page]),
            parser_version=PARSER_VERSION,
            metadata={"parser": "plain_text"},
        )

    def _parse_pdf(
        self,
        filename: str,
        file_extension: str,
        content: bytes,
    ) -> ParsedDocumentContent:
        import fitz  # type: ignore[import-untyped]

        document = fitz.open(stream=content, filetype="pdf")
        pages: list[ParsedPage] = []
        text_parts: list[str] = []
        cursor = 0
        for index, page in enumerate(document, start=1):
            page_text = normalize_extracted_text(page.get_text("text"))
            if text_parts:
                text_parts.append("\n\n")
                cursor += 2
            start = cursor
            text_parts.append(page_text)
            cursor += len(page_text)
            pages.append(
                ParsedPage(
                    page_number=index,
                    text=page_text,
                    char_start=start,
                    char_end=cursor,
                )
            )

        full_text = "".join(text_parts).strip()
        if not full_text:
            return self._parse_pdf_with_ocr(
                filename=filename,
                file_extension=file_extension,
                content=content,
                page_count=len(pages),
            )

        return ParsedDocumentContent(
            filename=filename,
            file_extension=file_extension,
            text=full_text,
            pages=pages,
            sections=detect_sections(full_text, pages=pages),
            parser_version=PARSER_VERSION,
            metadata={"parser": "pymupdf", "page_count": len(pages)},
        )

    def _parse_pdf_with_ocr(
        self,
        *,
        filename: str,
        file_extension: str,
        content: bytes,
        page_count: int,
    ) -> ParsedDocumentContent:
        if not self.settings.ocr_enabled:
            raise DocumentNeedsOcrError("PDF does not contain extractable text.")
        ocr_result = self.ocr_service.extract_text(content=content, page_count=page_count)
        normalized_text = normalize_extracted_text(ocr_result.text)
        if not normalized_text:
            raise DocumentNeedsOcrError("OCR completed but did not extract text.")
        page = ParsedPage(
            page_number=1,
            text=normalized_text,
            char_start=0,
            char_end=len(normalized_text),
        )
        return ParsedDocumentContent(
            filename=filename,
            file_extension=file_extension,
            text=normalized_text,
            pages=[page],
            sections=detect_sections(normalized_text, pages=[page]),
            parser_version=PARSER_VERSION,
            metadata={
                "parser": "ocrmypdf",
                "page_count": page_count,
                "ocr_fallback": True,
                **ocr_result.metadata,
            },
        )

    def _parse_docx(
        self,
        filename: str,
        file_extension: str,
        content: bytes,
    ) -> ParsedDocumentContent:
        from docx import Document

        document = Document(bytes_to_stream(content))
        parts: list[str] = []
        for paragraph in document.paragraphs:
            if paragraph.text.strip():
                parts.append(paragraph.text.strip())
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))

        return self.parse_plain_text(
            filename=filename,
            file_extension=file_extension,
            text="\n".join(parts),
        )

    def _parse_rtf(
        self,
        filename: str,
        file_extension: str,
        content: bytes,
    ) -> ParsedDocumentContent:
        from striprtf.striprtf import rtf_to_text

        text = rtf_to_text(self._decode_text(content))
        return self.parse_plain_text(
            filename=filename,
            file_extension=file_extension,
            text=text,
        )

    @staticmethod
    def _decode_text(content: bytes) -> str:
        return content.decode("utf-8-sig")


def bytes_to_stream(content: bytes):
    from io import BytesIO

    return BytesIO(content)


def normalize_extracted_text(text: str) -> str:
    lines = [re.sub(r"[ \t]+", " ", line).strip() for line in text.replace("\r", "\n").split("\n")]
    normalized = "\n".join(line for line in lines if line)
    return re.sub(r"\n{3,}", "\n\n", normalized).strip()


def detect_sections(text: str, *, pages: list[ParsedPage]) -> list[ParsedSection]:
    del pages
    sections: list[ParsedSection] = []
    matches = list(re.finditer(r"(?m)^(\d+(?:\.\d+)*\.?\s+\S.+)$", text))
    for index, match in enumerate(matches):
        end = matches[index + 1].start() if index + 1 < len(matches) else len(text)
        sections.append(
            ParsedSection(
                title=match.group(1).strip(),
                page_start=1,
                char_start=match.start(),
                char_end=end,
            )
        )
    return sections
