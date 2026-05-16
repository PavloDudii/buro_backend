import asyncio
import uuid
from collections.abc import Sequence
from io import BytesIO
from types import SimpleNamespace

import pytest
from openai.lib._pydantic import to_strict_json_schema
from sqlalchemy import func, select
from structlog.testing import capture_logs

from src.core.db.session import AsyncSessionLocal
from src.models.document import (
    DocumentChunk,
    DocumentExtractionItem,
    DocumentProcessingRun,
    DocumentProcessingStatus,
    ParsedDocument,
    UploadedDocument,
)
from src.models.user import User, UserRole
from src.services.blob_storage import StoredBlob
from src.services.document_chunking import DocumentChunker
from src.services.document_parsing import (
    DocumentNeedsOcrError,
    OcrTextResult,
    DocumentParser,
    UnsupportedDocumentTypeError,
)
from src.services.document_processing import (
    DocumentProcessingService,
    ExtractedItem,
    ExtractionResultSchema,
    OpenAIDocumentExtractionClient,
    normalize_fts_text,
    post_process_extracted_items,
)


class FakeBlobStorage:
    def __init__(self) -> None:
        self.objects: dict[str, bytes] = {}
        self.deleted_pathnames: list[str] = []

    async def put_document(
        self,
        *,
        user_id: uuid.UUID,
        document_id: uuid.UUID,
        safe_filename: str,
        content: bytes,
        content_type: str,
        uploaded_at,
    ) -> StoredBlob:
        del user_id, content_type, uploaded_at
        pathname = f"documents/{document_id}-{safe_filename}"
        self.objects[pathname] = content
        return StoredBlob(
            pathname=pathname,
            url=f"https://blob.test/{pathname}",
            download_url=f"https://blob.test/{pathname}?download=1",
        )

    async def get_document_content(self, pathname: str) -> bytes:
        return self.objects[pathname]

    async def delete_documents(self, pathnames: Sequence[str]) -> None:
        self.deleted_pathnames.extend(pathnames)


class FakeEmbeddingClient:
    def __init__(self, *, fail: bool = False) -> None:
        self.fail = fail
        self.inputs: list[str] = []

    async def embed_texts(self, texts: Sequence[str]) -> list[list[float]]:
        self.inputs.extend(texts)
        if self.fail:
            raise RuntimeError("embedding unavailable")
        return [[float(index + 1)] * 1536 for index, _ in enumerate(texts)]


class FakeExtractionClient:
    def __init__(self, *, fail: bool = False) -> None:
        self.fail = fail
        self.seen_text = ""

    async def extract_items(self, *, document_text: str, chunks: Sequence[DocumentChunk]):
        del chunks
        self.seen_text = document_text
        if self.fail:
            raise RuntimeError("extractor unavailable")
        return [
            ExtractedItem(
                type="person",
                value_json={"person_name": "Оскар Саєнко", "title": "Відповідальна особа"},
                confidence=0.91,
                source="openai",
                evidence_text="Оскар Саєнко відповідає за подання документів.",
                page_start=1,
                page_end=1,
                char_start=0,
                char_end=52,
            )
        ]


class BlockingExtractionClient:
    def __init__(self) -> None:
        self.started = asyncio.Event()
        self.release = asyncio.Event()

    async def extract_items(self, *, document_text: str, chunks: Sequence[DocumentChunk]):
        del document_text, chunks
        self.started.set()
        await self.release.wait()
        return [
            ExtractedItem(
                type="person",
                value_json={"person_name": "Live Metrics"},
                confidence=0.9,
                source="openai",
            )
        ]


class OcrEnabledSettings:
    ocr_enabled = True
    ocr_languages = "ukr+eng"
    ocr_timeout_seconds = 300
    ocr_max_pages = 80


class FakeOcrService:
    def __init__(self, text: str = "OCR text\n1. OCR section") -> None:
        self.text = text
        self.calls = 0

    def extract_text(self, *, content: bytes, page_count: int) -> OcrTextResult:
        del content
        self.calls += 1
        return OcrTextResult(
            text=self.text,
            metadata={
                "ocr_engine": "fake",
                "ocr_languages": "ukr+eng",
                "ocr_page_count": page_count,
                "ocr_duration_ms": 1,
                "ocr_timeout": False,
            },
        )


def pdf_with_text() -> bytes:
    import fitz  # type: ignore[import-untyped]

    document = fitz.open()
    page = document.new_page()
    page.insert_text((72, 72), "University rules\n1. General requirements")
    return document.tobytes()


def pdf_without_text() -> bytes:
    import fitz  # type: ignore[import-untyped]

    document = fitz.open()
    document.new_page()
    return document.tobytes()


def docx_with_text() -> bytes:
    from docx import Document

    document = Document()
    document.add_heading("Порядок подання документів", level=1)
    document.add_paragraph("Студент подає заяву до деканату.")
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


async def create_uploaded_document(
    *,
    storage: FakeBlobStorage,
    filename: str,
    content: bytes,
) -> uuid.UUID:
    async with AsyncSessionLocal() as session:
        user = User(
            email=f"{uuid.uuid4()}@example.com",
            full_name="Processing Admin",
            password_hash="hash",
            role=UserRole.ADMIN,
        )
        session.add(user)
        await session.flush()
        document_id = uuid.uuid4()
        storage_key = f"documents/{document_id}-{filename}"
        storage.objects[storage_key] = content
        document = UploadedDocument(
            id=document_id,
            original_filename=filename,
            safe_filename=filename,
            content_type="application/octet-stream",
            file_extension=filename.rsplit(".", 1)[-1],
            size_bytes=len(content),
            sha256_hash="a" * 64,
            storage_key=storage_key,
            uploaded_by_id=user.id,
            processing_status=DocumentProcessingStatus.QUEUED,
        )
        session.add(document)
        await session.commit()
        return document_id


@pytest.mark.anyio
async def test_parser_supports_text_document_formats() -> None:
    parser = DocumentParser()

    parsed_pdf = await parser.parse(
        filename="rules.pdf",
        file_extension="pdf",
        content=pdf_with_text(),
    )
    parsed_docx = await parser.parse(
        filename="procedure.docx",
        file_extension="docx",
        content=docx_with_text(),
    )
    parsed_txt = await parser.parse(
        filename="notes.txt",
        file_extension="txt",
        content="Контакт: деканат@example.edu.ua".encode(),
    )
    parsed_md = await parser.parse(
        filename="rules.md",
        file_extension="md",
        content="# Розділ\n\nСтудент подає заяву.".encode(),
    )
    parsed_rtf = await parser.parse(
        filename="rules.rtf",
        file_extension="rtf",
        content=b"{\\rtf1\\ansi \\u1055?\\u1088?\\u1072?\\u1074?\\u1080?\\u1083?\\u1072?}",
    )

    assert "University rules" in parsed_pdf.text
    assert parsed_pdf.pages[0].page_number == 1
    assert "Порядок подання" in parsed_docx.text
    assert "деканат@example.edu.ua" in parsed_txt.text
    assert "Студент подає заяву" in parsed_md.text
    assert parsed_rtf.text.strip()


@pytest.mark.anyio
async def test_parser_marks_scanned_pdf_and_unsupported_formats() -> None:
    parser = DocumentParser()

    with pytest.raises(DocumentNeedsOcrError):
        await parser.parse(filename="scan.pdf", file_extension="pdf", content=pdf_without_text())

    with pytest.raises(UnsupportedDocumentTypeError):
        await parser.parse(filename="table.csv", file_extension="csv", content=b"a,b\n1,2")


@pytest.mark.anyio
async def test_parser_uses_ocr_fallback_for_scanned_pdf_when_enabled() -> None:
    ocr = FakeOcrService()
    parser = DocumentParser(settings=OcrEnabledSettings(), ocr_service=ocr)

    parsed = await parser.parse(filename="scan.pdf", file_extension="pdf", content=pdf_without_text())

    assert "OCR text" in parsed.text
    assert parsed.metadata["ocr_fallback"] is True
    assert parsed.metadata["ocr_engine"] == "fake"
    assert ocr.calls == 1


def test_chunker_preserves_section_paths_and_source_spans() -> None:
    parsed = DocumentParser.parse_plain_text(
        filename="rules.md",
        file_extension="md",
        text=(
            "1. Загальні положення\n"
            "Студент має право на інформацію.\n\n"
            "2. Подання заяв\n"
            "Заява подається до деканату у письмовій формі."
        ),
    )

    chunks = DocumentChunker(target_tokens=8, overlap_tokens=2).chunk(parsed)

    assert len(chunks) >= 2
    assert chunks[0].section_path == "1. Загальні положення"
    assert chunks[-1].section_path == "2. Подання заяв"
    assert all(chunk.page_start == 1 and chunk.page_end == 1 for chunk in chunks)
    assert all(chunk.char_start < chunk.char_end for chunk in chunks)


def test_fts_normalization_is_ukrainian_safe() -> None:
    assert normalize_fts_text("  ЗАЯВУ\tпотрібно\nПОДАТИ  ") == "заяву потрібно подати"


def test_openai_extraction_schema_is_strict_output_compatible() -> None:
    schema = to_strict_json_schema(ExtractionResultSchema)

    assert "value_json" not in str(schema)
    assert_all_object_schemas_forbid_extra_properties(schema)


def test_extraction_guardrails_keep_only_people() -> None:
    items = post_process_extracted_items(
        [
            ExtractedItem(
                type="person",
                value_json={"person_name": "Оскар Саєнко", "title": "Відповідальна особа"},
                confidence=0.9,
                source="openai",
                evidence_text="Оскар Саєнко відповідає за подання документів.",
            ),
            ExtractedItem(
                type="department",
                value_json={"department": "ІКТА"},
                confidence=0.95,
                source="openai",
                evidence_text="ІКТА готує фахівців.",
            ),
            ExtractedItem(
                type="department",
                value_json={"department": "кафедра інтелектуальної мехатроніки і роботики (ІМР)"},
                confidence=0.9,
                source="openai",
                evidence_text="Кафедра інтелектуальної мехатроніки і роботики (ІМР) уклала меморандум.",
            ),
            ExtractedItem(
                type="form",
                value_json={
                    "title": "Бакалаврська спеціальність F7",
                    "form_name": "F7 Комп’ютерна інженерія (Комп’ютерна інженерія) (д)",
                    "raw_value": "F7 Комп’ютерна інженерія (Комп’ютерна інженерія) (д)",
                },
                confidence=0.9,
                source="openai",
                evidence_text="F7 Комп’ютерна інженерія (Комп’ютерна інженерія) (д)",
            ),
            ExtractedItem(
                type="contact",
                value_json={"email": "dekanat@example.edu.ua"},
                confidence=0.9,
                source="openai",
                evidence_text="Контакт: dekanat@example.edu.ua.",
            ),
            ExtractedItem(
                type="deadline",
                value_json={"date": "2026-06-01"},
                confidence=0.9,
                source="openai",
                evidence_text="Заяву необхідно подати до 1 червня 2026 року.",
            ),
            ExtractedItem(
                type="requirement",
                value_json={
                    "title": "Подання документів",
                    "requirement_text": "Необхідно подати заяву до деканату.",
                },
                confidence=0.9,
                source="openai",
                evidence_text="Необхідно подати заяву до деканату.",
            ),
            ExtractedItem(
                type="rule",
                value_json={"title": "Правило подання заяв"},
                confidence=0.9,
                source="openai",
                evidence_text="Заяви подаються у письмовій формі.",
            ),
        ]
    )

    assert [item.type for item in items] == ["person"]
    assert items[0].value_json["person_name"] == "Оскар Саєнко"


def test_extraction_guardrails_filter_institutions_and_departments() -> None:
    items = post_process_extracted_items(
        [
            ExtractedItem(
                type="department",
                value_json={
                    "title": "ІНСТИТУТ КОМП’ЮТЕРНИХ ТЕХНОЛОГІЙ, АВТОМАТИКИ ТА МЕТРОЛОГІЇ",
                    "department": "ІНСТИТУТ КОМП’ЮТЕРНИХ ТЕХНОЛОГІЙ, АВТОМАТИКИ ТА МЕТРОЛОГІЇ",
                },
                confidence=0.9,
                source="openai",
                evidence_text="Інститут готує фахівців.",
            ),
            ExtractedItem(
                type="department",
                value_json={
                    "title": "ІКТА",
                    "department": "ІКТА",
                },
                confidence=0.95,
                source="openai",
                evidence_text="Інститут комп’ютерних технологій, автоматики та метрології.",
            ),
            ExtractedItem(
                type="department",
                value_json={
                    "title": "Кафедра інтелектуальної мехатроніки і роботики (ІМР)",
                    "department": "кафедра інтелектуальної мехатроніки і роботики (ІМР)",
                },
                confidence=0.9,
                source="openai",
                evidence_text="Кафедра інтелектуальної мехатроніки і роботики (ІМР) уклала меморандум.",
            ),
            ExtractedItem(
                type="department",
                value_json={
                    "title": "Партнери за програмами подвійних дипломів",
                    "department": "Технічний університет Ільменау (Німеччина); Університет дю Мен (Франція)",
                },
                confidence=0.9,
                source="openai",
                evidence_text="Технічний університет Ільменау; Університет дю Мен.",
            ),
        ]
    )

    assert items == []


@pytest.mark.anyio
async def test_openai_extraction_sends_chunk_scoped_prompts_without_full_document_duplication() -> None:
    class FakeResponses:
        def __init__(self) -> None:
            self.prompts: list[str] = []

        async def parse(self, *, model, input, text_format):
            del model, text_format
            self.prompts.append(input[-1]["content"])
            return SimpleNamespace(output_parsed=ExtractionResultSchema(items=[]))

    class FakeOpenAIClient:
        def __init__(self) -> None:
            self.responses = FakeResponses()

    fake_client = FakeOpenAIClient()
    extraction_client = OpenAIDocumentExtractionClient(
        SimpleNamespace(
            openai_api_key="test",
            openai_extraction_model="test-model",
            openai_extraction_concurrency=2,
        )
    )
    extraction_client._client = fake_client
    document_id = uuid.uuid4()
    chunks = [
        DocumentChunk(
            id=uuid.uuid4(),
            document_id=document_id,
            chunk_index=0,
            content="First chunk only.",
            fts_text="first chunk only",
            char_start=0,
            char_end=17,
        ),
        DocumentChunk(
            id=uuid.uuid4(),
            document_id=document_id,
            chunk_index=1,
            content="Second chunk only.",
            fts_text="second chunk only",
            char_start=18,
            char_end=36,
        ),
    ]

    items = await extraction_client.extract_items(
        document_text="FULL DOCUMENT TEXT SHOULD NOT BE SENT",
        chunks=chunks,
    )

    assert items == []
    assert len(fake_client.responses.prompts) == 2
    assert "FULL DOCUMENT TEXT SHOULD NOT BE SENT" not in str(fake_client.responses.prompts)
    assert fake_client.responses.prompts[0].count("First chunk only.") == 1
    assert fake_client.responses.prompts[1].count("Second chunk only.") == 1


def assert_all_object_schemas_forbid_extra_properties(schema: object) -> None:
    if isinstance(schema, dict):
        if schema.get("type") == "object":
            assert schema.get("additionalProperties") is False
        for value in schema.values():
            assert_all_object_schemas_forbid_extra_properties(value)
    elif isinstance(schema, list):
        for item in schema:
            assert_all_object_schemas_forbid_extra_properties(item)


@pytest.mark.anyio
async def test_processing_persists_parsed_text_chunks_embeddings_and_extractions() -> None:
    storage = FakeBlobStorage()
    document_id = await create_uploaded_document(
        storage=storage,
        filename="rules.txt",
        content=(
            "1. Загальні положення\n"
            "Заяву необхідно подати до 1 червня 2026 року.\n"
            "Контакт: dekanat@example.edu.ua"
        ).encode(),
    )

    with capture_logs() as logs:
        async with AsyncSessionLocal() as session:
            service = DocumentProcessingService(
                session,
                blob_storage=storage,
                embedding_client=FakeEmbeddingClient(),
                extraction_client=FakeExtractionClient(),
            )
            await service.process_document(document_id=document_id)

    async with AsyncSessionLocal() as session:
        document = await session.get(UploadedDocument, document_id)
        parsed = (
            await session.execute(
                select(ParsedDocument).where(ParsedDocument.document_id == document_id)
            )
        ).scalar_one()
        chunks = (
            await session.execute(select(DocumentChunk).where(DocumentChunk.document_id == document_id))
        ).scalars().all()
        extraction = (
            await session.execute(
                select(DocumentExtractionItem).where(
                    DocumentExtractionItem.document_id == document_id
                )
            )
        ).scalar_one()
        processing_run = (
            await session.execute(
                select(DocumentProcessingRun).where(
                    DocumentProcessingRun.document_id == document_id
                )
            )
        ).scalar_one()
        search_vector = (
            await session.execute(
                select(func.to_tsvector("simple", DocumentChunk.fts_text)).where(
                    DocumentChunk.document_id == document_id
                )
            )
        ).scalar_one()

    assert document is not None
    assert document.processing_status == DocumentProcessingStatus.COMPLETED
    assert document.processing_error is None
    assert document.processing_error_code is None
    assert document.processing_error_stage is None
    assert parsed.language == "ukrainian"
    assert "Заяву необхідно" in parsed.raw_text
    assert chunks
    assert all(chunk.embedding is not None and len(chunk.embedding) == 1536 for chunk in chunks)
    assert all(chunk.embedding_model == "text-embedding-3-small" for chunk in chunks)
    assert chunks[0].fts_text == normalize_fts_text(chunks[0].content)
    assert str(search_vector)
    assert extraction.type == "person"
    assert extraction.value_json["person_name"] == "Оскар Саєнко"
    assert extraction.evidence_text
    assert processing_run.status == DocumentProcessingStatus.COMPLETED
    assert processing_run.completed_at is not None
    assert processing_run.total_duration_ms is not None
    assert processing_run.total_duration_ms >= 0
    assert processing_run.error_code is None
    assert processing_run.error_stage is None
    assert set(processing_run.stage_metrics_json) == {
        "blob_read",
        "parse",
        "chunk",
        "embed",
        "extract",
        "persist",
    }
    assert all(
        metric["status"] == "completed"
        and isinstance(metric["duration_ms"], int)
        and metric["duration_ms"] >= 0
        and metric["started_at"]
        and metric["completed_at"]
        for metric in processing_run.stage_metrics_json.values()
    )
    assert processing_run.stage_metrics_json["blob_read"]["bytes_read"] == document.size_bytes
    assert processing_run.stage_metrics_json["parse"]["page_count"] == 1
    assert processing_run.stage_metrics_json["parse"]["section_count"] >= 1
    assert processing_run.stage_metrics_json["parse"]["text_char_count"] == len(parsed.raw_text)
    assert processing_run.stage_metrics_json["chunk"]["chunk_count"] == len(chunks)
    assert processing_run.stage_metrics_json["chunk"]["token_count_total"] >= len(chunks)
    assert processing_run.stage_metrics_json["embed"]["embedding_count"] == len(chunks)
    assert (
        processing_run.stage_metrics_json["embed"]["embedding_model"]
        == "text-embedding-3-small"
    )
    assert processing_run.stage_metrics_json["extract"]["extraction_count"] == 1
    assert processing_run.summary_metrics_json["chunk_count"] == len(chunks)
    assert processing_run.summary_metrics_json["embedding_count"] == len(chunks)
    assert processing_run.summary_metrics_json["extraction_count"] == 1
    assert "Заяву необхідно" not in str(processing_run.stage_metrics_json)
    assert "dekanat@example.edu.ua" not in str(processing_run.stage_metrics_json)
    assert "Заяву необхідно" not in str(logs)
    assert "dekanat@example.edu.ua" not in str(logs)
    assert {
        ("document_processing.stage.started", "blob_read"),
        ("document_processing.stage.completed", "blob_read"),
        ("document_processing.stage.started", "parse"),
        ("document_processing.stage.completed", "parse"),
        ("document_processing.stage.started", "chunk"),
        ("document_processing.stage.completed", "chunk"),
        ("document_processing.stage.started", "embed"),
        ("document_processing.stage.completed", "embed"),
        ("document_processing.stage.started", "extract"),
        ("document_processing.stage.completed", "extract"),
        ("document_processing.stage.started", "persist"),
        ("document_processing.stage.completed", "persist"),
        ("document_processing.completed", "complete"),
    }.issubset({(entry["event"], entry.get("stage")) for entry in logs})
    assert all(entry.get("document_id") == str(document_id) for entry in logs)
    assert all(entry.get("processing_run_id") == str(processing_run.id) for entry in logs)
    assert all(
        isinstance(entry.get("duration_ms"), int)
        for entry in logs
        if entry["event"] == "document_processing.stage.completed"
    )
    final_log = next(entry for entry in logs if entry["event"] == "document_processing.completed")
    assert final_log["total_duration_ms"] == processing_run.total_duration_ms


@pytest.mark.anyio
async def test_processing_persists_stage_metrics_while_run_is_active() -> None:
    storage = FakeBlobStorage()
    extraction_client = BlockingExtractionClient()
    document_id = await create_uploaded_document(
        storage=storage,
        filename="rules.txt",
        content=(
            "1. Загальні положення\n"
            "Заяву необхідно подати до 1 червня 2026 року.\n"
            "Контакт: dekanat@example.edu.ua"
        ).encode(),
    )

    async with AsyncSessionLocal() as session:
        service = DocumentProcessingService(
            session,
            blob_storage=storage,
            embedding_client=FakeEmbeddingClient(),
            extraction_client=extraction_client,
        )
        task = asyncio.create_task(service.process_document(document_id=document_id))
        await asyncio.wait_for(extraction_client.started.wait(), timeout=2)

        async with AsyncSessionLocal() as observer_session:
            processing_run = (
                await observer_session.execute(
                    select(DocumentProcessingRun).where(
                        DocumentProcessingRun.document_id == document_id
                    )
                )
            ).scalar_one()

        assert processing_run.status == DocumentProcessingStatus.PROCESSING
        assert processing_run.completed_at is None
        assert processing_run.stage_metrics_json["blob_read"]["status"] == "completed"
        assert processing_run.stage_metrics_json["parse"]["status"] == "completed"
        assert processing_run.stage_metrics_json["chunk"]["status"] == "completed"
        assert processing_run.stage_metrics_json["embed"]["status"] == "completed"
        assert processing_run.stage_metrics_json["extract"]["status"] == "started"
        assert processing_run.stage_metrics_json["extract"]["started_at"]

        extraction_client.release.set()
        await task


@pytest.mark.anyio
async def test_processing_classifies_embedding_failure_with_safe_error_fields() -> None:
    storage = FakeBlobStorage()
    document_id = await create_uploaded_document(
        storage=storage,
        filename="rules.txt",
        content="1. Загальні положення\nТекст документа.".encode(),
    )

    with capture_logs() as logs:
        async with AsyncSessionLocal() as session:
            await DocumentProcessingService(
                session,
                blob_storage=storage,
                embedding_client=FakeEmbeddingClient(fail=True),
                extraction_client=FakeExtractionClient(),
            ).process_document(document_id=document_id)

    async with AsyncSessionLocal() as session:
        document = await session.get(UploadedDocument, document_id)
        processing_run = (
            await session.execute(
                select(DocumentProcessingRun).where(
                    DocumentProcessingRun.document_id == document_id
                )
            )
        ).scalar_one()

    assert document is not None
    assert document.processing_status == DocumentProcessingStatus.FAILED
    assert document.processing_error_code == "embedding_failed"
    assert document.processing_error_stage == "embed"
    assert document.processing_error == "Document embedding failed."
    assert processing_run.status == DocumentProcessingStatus.FAILED
    assert processing_run.completed_at is not None
    assert processing_run.total_duration_ms is not None
    assert processing_run.error_code == "embedding_failed"
    assert processing_run.error_stage == "embed"
    assert processing_run.error_message == "Document embedding failed."
    assert processing_run.stage_metrics_json["blob_read"]["status"] == "completed"
    assert processing_run.stage_metrics_json["parse"]["status"] == "completed"
    assert processing_run.stage_metrics_json["chunk"]["status"] == "completed"
    assert processing_run.stage_metrics_json["embed"]["status"] == "failed"
    assert processing_run.stage_metrics_json["embed"]["error_code"] == "embedding_failed"
    assert processing_run.stage_metrics_json["embed"]["error_type"] == "RuntimeError"
    assert isinstance(processing_run.stage_metrics_json["embed"]["duration_ms"], int)
    assert "Текст документа" not in str(processing_run.stage_metrics_json)
    assert any(
        entry["event"] == "document_processing.stage.failed"
        and entry["stage"] == "embed"
        and entry["error_code"] == "embedding_failed"
        and entry["error_type"] == "RuntimeError"
        and isinstance(entry["duration_ms"], int)
        for entry in logs
    )
    final_log = next(entry for entry in logs if entry["event"] == "document_processing.failed")
    assert final_log["total_duration_ms"] == processing_run.total_duration_ms
    assert "Текст документа" not in str(logs)


@pytest.mark.anyio
async def test_processing_marks_unsupported_needs_ocr_and_failed_statuses() -> None:
    unsupported_storage = FakeBlobStorage()
    unsupported_id = await create_uploaded_document(
        storage=unsupported_storage,
        filename="table.csv",
        content=b"a,b\n1,2",
    )
    ocr_storage = FakeBlobStorage()
    ocr_id = await create_uploaded_document(
        storage=ocr_storage,
        filename="scan.pdf",
        content=pdf_without_text(),
    )
    failed_storage = FakeBlobStorage()
    failed_id = await create_uploaded_document(
        storage=failed_storage,
        filename="rules.txt",
        content=b"\xd0\x9f\xd1\x80\xd0\xb0\xd0\xb2\xd0\xb8\xd0\xbb\xd0\xb0",
    )

    async with AsyncSessionLocal() as session:
        await DocumentProcessingService(
            session,
            blob_storage=unsupported_storage,
            embedding_client=FakeEmbeddingClient(),
            extraction_client=FakeExtractionClient(),
        ).process_document(document_id=unsupported_id)
        await DocumentProcessingService(
            session,
            blob_storage=ocr_storage,
            embedding_client=FakeEmbeddingClient(),
            extraction_client=FakeExtractionClient(),
        ).process_document(document_id=ocr_id)
        await DocumentProcessingService(
            session,
            blob_storage=failed_storage,
            embedding_client=FakeEmbeddingClient(),
            extraction_client=FakeExtractionClient(fail=True),
        ).process_document(document_id=failed_id)

    async with AsyncSessionLocal() as session:
        unsupported = await session.get(UploadedDocument, unsupported_id)
        ocr = await session.get(UploadedDocument, ocr_id)
        failed = await session.get(UploadedDocument, failed_id)
        failed_chunks = (
            await session.execute(select(DocumentChunk).where(DocumentChunk.document_id == failed_id))
        ).scalars().all()

    assert unsupported is not None
    assert ocr is not None
    assert failed is not None
    assert unsupported.processing_status == DocumentProcessingStatus.UNSUPPORTED
    assert unsupported.processing_error_code == "unsupported_file_type"
    assert unsupported.processing_error_stage == "parse"
    assert ocr.processing_status == DocumentProcessingStatus.NEEDS_OCR
    assert ocr.processing_error_code == "needs_ocr"
    assert ocr.processing_error_stage == "parse"
    assert failed.processing_status == DocumentProcessingStatus.FAILED
    assert failed.processing_error_code == "extraction_provider_failed"
    assert failed.processing_error_stage == "extract"
    assert failed.processing_error == "Structured extraction failed."
    assert failed_chunks == []


@pytest.mark.anyio
async def test_processing_ocr_fallback_reaches_completed_status() -> None:
    storage = FakeBlobStorage()
    document_id = await create_uploaded_document(
        storage=storage,
        filename="scan.pdf",
        content=pdf_without_text(),
    )
    parser = DocumentParser(settings=OcrEnabledSettings(), ocr_service=FakeOcrService())

    async with AsyncSessionLocal() as session:
        await DocumentProcessingService(
            session,
            blob_storage=storage,
            embedding_client=FakeEmbeddingClient(),
            extraction_client=FakeExtractionClient(),
            parser=parser,
        ).process_document(document_id=document_id)

    async with AsyncSessionLocal() as session:
        document = await session.get(UploadedDocument, document_id)
        run = (
            await session.execute(select(DocumentProcessingRun).where(DocumentProcessingRun.document_id == document_id))
        ).scalar_one()

    assert document is not None
    assert document.processing_status == DocumentProcessingStatus.COMPLETED
    assert run.stage_metrics_json["parse"]["ocr_engine"] == "fake"
    assert run.summary_metrics_json["ocr_engine"] == "fake"


@pytest.mark.anyio
async def test_reprocessing_replaces_existing_artifacts_idempotently() -> None:
    storage = FakeBlobStorage()
    document_id = await create_uploaded_document(
        storage=storage,
        filename="rules.txt",
        content=b"\xd0\x9f\xd0\xb5\xd1\x80\xd1\x88\xd0\xb0 \xd0\xb2\xd0\xb5\xd1\x80\xd1\x81\xd1\x96\xd1\x8f",
    )

    async with AsyncSessionLocal() as session:
        service = DocumentProcessingService(
            session,
            blob_storage=storage,
            embedding_client=FakeEmbeddingClient(),
            extraction_client=FakeExtractionClient(),
        )
        await service.process_document(document_id=document_id)

    storage.objects[f"documents/{document_id}-rules.txt"] = (
        "2. Нова версія\nОновлений порядок подання заяв."
    ).encode()

    async with AsyncSessionLocal() as session:
        service = DocumentProcessingService(
            session,
            blob_storage=storage,
            embedding_client=FakeEmbeddingClient(),
            extraction_client=FakeExtractionClient(),
        )
        await service.queue_document(document_id=document_id)
        await service.process_document(document_id=document_id)

    async with AsyncSessionLocal() as session:
        parsed_rows = (
            await session.execute(
                select(ParsedDocument).where(ParsedDocument.document_id == document_id)
            )
        ).scalars().all()
        chunks = (
            await session.execute(select(DocumentChunk).where(DocumentChunk.document_id == document_id))
        ).scalars().all()
        processing_runs = (
            await session.execute(
                select(DocumentProcessingRun)
                .where(DocumentProcessingRun.document_id == document_id)
                .order_by(DocumentProcessingRun.started_at)
            )
        ).scalars().all()

    assert len(parsed_rows) == 1
    assert "Оновлений порядок" in parsed_rows[0].raw_text
    assert chunks
    assert all("Перша версія" not in chunk.content for chunk in chunks)
    assert len(processing_runs) == 2
    assert len({run.id for run in processing_runs}) == 2
    assert all(run.status == DocumentProcessingStatus.COMPLETED for run in processing_runs)
